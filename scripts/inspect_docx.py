"""Diagnostic: print the structural shape of a .docx file.

Usage:
    python -m scripts.inspect_docx "data/faq_raw/Newlife_Medicare_Comprehensive_AI_FAQ.docx"

Prints:
  - count of paragraphs by style
  - whether tables exist and their shapes
  - regex-marker hit counts (Q:/A:/numbered/etc.)
  - first few examples of each detected pattern

This script DOES NOT extract anything. It only diagnoses so we can pick the
right extraction strategy.
"""

from __future__ import annotations

import sys
from collections import Counter
from pathlib import Path

from docx import Document

from scripts._markers import Q_MARKERS, A_MARKERS, is_question_line, is_answer_line


def inspect(path: Path) -> None:
    print(f"\n=== INSPECT DOCX: {path.name} ===")
    if not path.exists():
        print(f"  ERROR: file not found: {path}")
        sys.exit(1)

    doc = Document(str(path))

    style_counts: Counter[str] = Counter()
    bold_q_count = 0
    ends_with_qmark = 0
    nonempty_paragraphs = 0
    marker_hits: Counter[str] = Counter()
    samples: dict[str, list[str]] = {"Q_marker": [], "A_marker": [], "bold_q": [], "heading": []}

    for para in doc.paragraphs:
        text = (para.text or "").strip()
        style_name = para.style.name if para.style else "<no style>"
        if text:
            nonempty_paragraphs += 1
            style_counts[style_name] += 1

            if is_question_line(text):
                marker_hits["Q_marker"] += 1
                if len(samples["Q_marker"]) < 3:
                    samples["Q_marker"].append(text[:120])
            if is_answer_line(text):
                marker_hits["A_marker"] += 1
                if len(samples["A_marker"]) < 3:
                    samples["A_marker"].append(text[:120])

            if text.endswith("?"):
                ends_with_qmark += 1
                all_bold = bool(para.runs) and all(
                    (r.bold or (r.font and r.font.bold)) for r in para.runs if r.text.strip()
                )
                if all_bold:
                    bold_q_count += 1
                    if len(samples["bold_q"]) < 3:
                        samples["bold_q"].append(text[:120])

            if style_name.startswith("Heading"):
                if len(samples["heading"]) < 3:
                    samples["heading"].append(f"[{style_name}] {text[:100]}")

    table_shapes: list[tuple[int, int]] = []
    table_first_rows: list[list[str]] = []
    for tbl in doc.tables:
        rows = len(tbl.rows)
        cols = len(tbl.columns) if rows else 0
        table_shapes.append((rows, cols))
        if tbl.rows:
            first = [cell.text.strip()[:60] for cell in tbl.rows[0].cells]
            table_first_rows.append(first)

    print(f"\nParagraphs (non-empty): {nonempty_paragraphs}")
    print(f"Paragraphs ending with '?': {ends_with_qmark}  (of which fully bold: {bold_q_count})")

    print("\nParagraph styles used:")
    for style, count in style_counts.most_common():
        print(f"  {count:>4}  {style}")

    print("\nMarker regex hits:")
    print(f"  Q-style markers: {marker_hits.get('Q_marker', 0)}")
    print(f"  A-style markers: {marker_hits.get('A_marker', 0)}")
    print(f"  Bold-question pattern (bold paragraph ending in '?'): {bold_q_count}")

    print(f"\nTables found: {len(table_shapes)}")
    for i, (r, c) in enumerate(table_shapes[:5]):
        first = table_first_rows[i] if i < len(table_first_rows) else []
        print(f"  table[{i}]: {r} rows x {c} cols   first row: {first}")
    if len(table_shapes) > 5:
        print(f"  ... and {len(table_shapes) - 5} more tables")

    print("\nSample lines per detected pattern:")
    for key, items in samples.items():
        if items:
            print(f"  [{key}]")
            for s in items:
                print(f"    - {s}")

    print("\n--- STRATEGY RECOMMENDATION ---")
    recs: list[tuple[str, int]] = []
    if marker_hits.get("Q_marker", 0) >= 3:
        recs.append(("markers", marker_hits["Q_marker"]))
    if bold_q_count >= 3:
        recs.append(("bold_q", bold_q_count))
    heading_q_count = sum(
        1 for p in doc.paragraphs
        if (p.style.name if p.style else "").startswith("Heading")
        and (p.text or "").strip().endswith("?")
    )
    if heading_q_count >= 3:
        recs.append(("headings", heading_q_count))
    if any(c >= 2 for _, c in table_shapes):
        recs.append(("tables", sum(r for r, c in table_shapes if c >= 2)))

    if not recs:
        print("  No strong pattern detected. Manual review needed.")
    else:
        recs.sort(key=lambda x: -x[1])
        print(f"  Recommended primary strategy: {recs[0][0]}  (evidence: {recs[0][1]} hits)")
        if len(recs) > 1:
            print(f"  Secondary strategies: {[r[0] for r in recs[1:]]}")


def main() -> None:
    if len(sys.argv) < 2:
        print("Usage: python -m scripts.inspect_docx <path-to-docx> [more paths...]")
        sys.exit(2)
    for arg in sys.argv[1:]:
        inspect(Path(arg))


if __name__ == "__main__":
    main()
